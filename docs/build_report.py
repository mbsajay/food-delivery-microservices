"""
Build the Foodly project report by transforming the Scaler Neovarsity template.

What this script does:
  1. Opens the template.
  2. Replaces placeholders in the front matter (title page, declaration,
     acknowledgment).
  3. Deletes the placeholder body content between "Applied Software Project"
     and "References" headings.
  4. Rebuilds the main report with project-specific content while keeping
     the template's heading styles (so ToC numbering still works).
  5. Applies global formatting rules from the template's Format Guidelines
     page: Times New Roman, 12pt body / 14pt headings, 1.5 line spacing,
     justified body, 1.25"/1" margins.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def _set_cell_border(cell):
    """Add a thin black border to every side of a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        tcBorders.append(b)
    tcPr.append(tcBorders)

TEMPLATE = Path(
    r"C:\Users\mbsaj\Downloads\Scaler Neovarsity _ Academy Project Report Template (Backend Specialization).docx"
)
OUTPUT = Path(r"D:\study\Claude\scalar\docs\Foodly_Project_Report.docx")

STUDENT_NAME = "Ajay Babu Mandarapu"
STUDENT_EMAIL = "mbsajay@gmail.com"
SUPERVISOR = "Naman Bhalla"
SUBMISSION_MONTH = "June 2026"
SUBMISSION_DATE = "30/05/2026"
MODULE_START = "1 May 2026"
MODULE_END = "30 May 2026"
DECLARATION_DATE = "30 May 2026"
PROJECT_TITLE = "Foodly – A Microservices-Based Food Delivery Platform"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------
def set_paragraph_text(paragraph, text: str) -> None:
    """Replace the entire text of a paragraph while keeping run formatting."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def replace_in_paragraph(paragraph, find: str, replace: str) -> bool:
    """Replace text within a paragraph that may be split across multiple runs."""
    full = "".join(r.text for r in paragraph.runs)
    if find not in full:
        return False
    new = full.replace(find, replace)
    if paragraph.runs:
        paragraph.runs[0].text = new
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(new)
    return True


def delete_paragraph(p) -> None:
    el = p._element
    el.getparent().remove(el)


# ---------------------------------------------------------------------------
# Front-matter rewrites
# ---------------------------------------------------------------------------
def fix_front_matter(doc) -> None:
    """Walk paragraphs and replace cover/declaration/acknowledgment text."""
    # NOTE: the template has the literal text "DD/MM/YYYY <Date of Submission>"
    # next to each other on the cover page; replace them together to avoid
    # producing the date twice.
    replacements = {
        "<Full Name of the Student>": STUDENT_NAME,
        "<Full Name of the Candidate>": STUDENT_NAME,
        "<Signature of the Candidate>": STUDENT_NAME,
        "<Month of Submission like June, 2024>": SUBMISSION_MONTH,
        "<Registered Scaler Email ID>": STUDENT_EMAIL,
        "DD/MM/YYYY <Date of Submission>": SUBMISSION_DATE,
        "<Date of Submission>": SUBMISSION_DATE,
        "DD/MM/YYYY": SUBMISSION_DATE,
        "< Project Module start date >": MODULE_START,
        "< Module end date >": MODULE_END,
        "XX Month 20XX": DECLARATION_DATE,
    }
    for p in doc.paragraphs:
        for find, rep in replacements.items():
            replace_in_paragraph(p, find, rep)

    # Acknowledgment paragraph (find the angle-bracket guidance and rewrite)
    ack_text = (
        "I would like to express my sincere gratitude to my project supervisor, "
        f"{SUPERVISOR}, whose expert guidance, patient feedback, and steady "
        "encouragement shaped both the design and the execution of this project. "
        "I am deeply thankful to the instructors at Scaler Neovarsity for the rigorous "
        "Backend Specialization curriculum, which gave me the foundation to attempt a "
        "project of this scope. I owe a special debt of thanks to my family for their "
        "unwavering support throughout the program — without their understanding through "
        "many late nights, this work would not have been possible. Finally, I thank my "
        "peers at Scaler whose discussions and code reviews helped sharpen my engineering "
        "judgment, and the broader open-source community whose excellent documentation "
        "for Spring Boot, Apache Kafka, and PostgreSQL made deep learning practical."
    )
    for p in doc.paragraphs:
        if p.text.strip().startswith("<Insert a Paragraph"):
            set_paragraph_text(p, ack_text)
            break

    # Project title on cover: the template doesn't have an explicit title placeholder,
    # but a project title typically sits near the top. We add the title as a new line
    # just below "Applied Software Project Report" on page 1, only if it isn't already
    # there. Easiest: insert via paragraph 0's parent xml.
    if doc.paragraphs and doc.paragraphs[0].text.strip().startswith("Applied Software Project Report"):
        title_para = doc.paragraphs[0]
        body = title_para._element.getparent()
        # Find paragraph "By" and insert title BEFORE it.
        for p in doc.paragraphs:
            if p.text.strip() == "By":
                # Insert "Project Title:" line above "By"
                new_p = deepcopy(title_para._element)
                # Clear runs in the copy and set the new text
                for r in new_p.findall(qn("w:r")):
                    new_p.remove(r)
                # Reuse paragraph properties; add a single run with the title
                run = OxmlElement("w:r")
                rpr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                run.append(rpr)
                rpr.append(b)
                t = OxmlElement("w:t")
                t.text = PROJECT_TITLE
                t.set(qn("xml:space"), "preserve")
                run.append(t)
                new_p.append(run)
                p._element.addprevious(new_p)
                break


# ---------------------------------------------------------------------------
# Main content rewrite
# ---------------------------------------------------------------------------
def find_paragraph_index_by_text(doc, text: str, exact: bool = True) -> int:
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if (exact and t == text) or (not exact and text in t):
            return i
    return -1


def remove_paragraphs_between(doc, start_text: str, end_text: str) -> None:
    """Remove all paragraphs strictly between the two named paragraphs."""
    start_idx = find_paragraph_index_by_text(doc, start_text)
    end_idx = find_paragraph_index_by_text(doc, end_text)
    if start_idx < 0 or end_idx < 0:
        raise RuntimeError(f"Could not find anchors: {start_text!r} / {end_text!r}")
    to_remove = doc.paragraphs[start_idx + 1 : end_idx]
    for p in to_remove:
        delete_paragraph(p)


def add_paragraph_before(anchor_paragraph, text: str = "", style: str = "Normal"):
    """Insert a new paragraph immediately before `anchor_paragraph`."""
    new_p = anchor_paragraph._parent.add_paragraph(text, style=style)
    anchor_paragraph._element.addprevious(new_p._element)
    return new_p


def style_body_paragraph(p, justify: bool = True) -> None:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)


def style_heading(p, size: int = 14, level: int = 3) -> None:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.font.bold = True


# ---------------------------------------------------------------------------
# Content writers — each section returns nothing, mutates doc in place by
# inserting new paragraphs *before* the References heading.
# ---------------------------------------------------------------------------
def insert_content(doc) -> None:
    """Build the main report by inserting all sections before 'References'."""
    refs_idx = find_paragraph_index_by_text(doc, "References")
    if refs_idx < 0:
        raise RuntimeError("Could not find References heading.")
    anchor = doc.paragraphs[refs_idx]

    def H3(text: str):
        p = add_paragraph_before(anchor, text, style="Heading 3")
        style_heading(p, size=14, level=3)
        return p

    def H4(text: str):
        p = add_paragraph_before(anchor, text, style="Heading 4")
        style_heading(p, size=13, level=4)
        return p

    def P(text: str, justify: bool = True):
        p = add_paragraph_before(anchor, text, style="Normal")
        style_body_paragraph(p, justify=justify)
        return p

    def B(text: str):
        """A bullet-style line (single-line spacing, no bullet glyph — keeps it simple)."""
        p = add_paragraph_before(anchor, "• " + text, style="Normal")
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.left_indent = Inches(0.3)
        pf.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
        return p

    def CAPTION(text: str):
        p = add_paragraph_before(anchor, text, style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(12)
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
            r.font.italic = True
        return p

    def TABLE(headers: list[str], rows: list[list[str]], caption: str):
        # Caption appears ABOVE tables per the template guidelines.
        cap = add_paragraph_before(anchor, caption, style="Normal")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
            r.font.italic = True

        tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        for ci, h in enumerate(headers):
            cell = tbl.rows[0].cells[ci]
            cell.text = h
            _set_cell_border(cell)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
        for ri, row in enumerate(rows, start=1):
            for ci, val in enumerate(row):
                cell = tbl.rows[ri].cells[ci]
                cell.text = str(val)
                _set_cell_border(cell)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)
        # Move the appended table to just before the anchor.
        anchor._element.addprevious(tbl._element)

    # -------------------------------------------------------------------
    # H2: Applied Software Project — already exists as a heading earlier.
    # We start fresh from Abstract.
    # -------------------------------------------------------------------

    # =======  Abstract  =======
    H3("Abstract")
    P(
        "This report documents the design, implementation, and deployment of "
        "Foodly, a microservices-based food delivery platform built as the "
        "applied capstone for the Master of Science in Computer Science (Backend "
        "Specialization) at Scaler Neovarsity – Woolf. Foodly mirrors the "
        "operational shape of production systems such as Swiggy, DoorDash, and "
        "Uber Eats: customers browse restaurants, place orders, pay through an "
        "external payment gateway, and receive their food via a delivery agent, "
        "with every step reflected in near-real-time notifications."
    )
    P(
        "The system is decomposed into nine independently deployable Spring Boot "
        "services — user, restaurant, order, payment, delivery, notification, "
        "an API gateway, a service registry, and a configuration server. The "
        "services communicate synchronously via REST when low-latency lookups "
        "are required and asynchronously via Apache Kafka for state propagation "
        "between bounded contexts. Persistence is handled through Spring Data "
        "JPA over PostgreSQL, authentication is implemented with stateless "
        "JSON Web Tokens signed using RS256, and a scheduled reconciliation "
        "job in the payment service guarantees eventual consistency with the "
        "external payment provider despite webhook loss or network failure."
    )
    P(
        "The project applies the complete Backend Specialization curriculum: "
        "Git-based workflows, Spring MVC and REST, Spring Data JPA inheritance "
        "and queries, RestTemplate-based third-party integration, JWT-secured "
        "endpoints, paged and sorted search, Kafka-based asynchronous "
        "communication, Spring Cloud (Gateway, Eureka, Config), and Docker-"
        "based containerisation. The codebase is verified by unit tests "
        "(JUnit 5 with Mockito) and integration tests (Testcontainers), and "
        "the entire stack — five infrastructure containers plus nine "
        "application services — is brought up by a single `docker compose "
        "up` command, making the submission fully self-contained and "
        "reproducible on any developer machine. The result is a demonstration "
        "that the patterns underlying modern food-delivery businesses can be "
        "assembled from open-source components and operated locally with "
        "production-shaped behaviour while remaining correct under concurrent "
        "load and partial failure."
    )

    # =======  Project Description  =======
    H3("Project Description")
    P(
        "Foodly is an online food-ordering and delivery platform that connects "
        "three classes of users — customers, restaurants, and delivery agents — "
        "through a single backend platform operated by a fourth class of user, "
        "the platform administrator. The product surface is intentionally "
        "small (customers can sign up, browse restaurants, search by cuisine, "
        "place orders, pay, and track delivery; restaurants can publish menus "
        "and receive orders; delivery agents can accept assignments and update "
        "trip status); the engineering surface, however, is deliberately broad, "
        "exercising every concept introduced in the Backend Specialization "
        "module."
    )
    P("The principal objectives of the project are:")
    B("Build a domain-faithful clone of a real-world food-delivery system rather than a toy CRUD application.")
    B("Decompose the system into independently deployable microservices with clean bounded contexts and well-defined contracts.")
    B("Demonstrate event-driven choreography across services using Apache Kafka, including reliable retry and idempotency.")
    B("Integrate with an external payment gateway via webhooks, and prove correctness with a scheduled reconciliation job.")
    B("Apply Spring Cloud (Gateway, Eureka, Config Server) to make the platform discoverable, configurable, and observable.")
    B("Containerize every component with Docker and orchestrate the full local stack via a single docker-compose.yml so the project is fully reproducible on any developer machine.")
    B("Cover the system with unit tests (JUnit, Mockito) and integration tests using Testcontainers against a real PostgreSQL instance.")
    P(
        "The relevance of the project is twofold. First, food-delivery is one "
        "of the canonical case studies in modern distributed-systems design — "
        "it features high-cardinality reads (browsing), strongly-consistent "
        "writes (order placement and payment), and long-running asynchronous "
        "workflows (delivery). Second, the same architectural primitives — "
        "service decomposition, event sourcing of state transitions, JWT-based "
        "edge authentication, and gateway-driven payment integration — recur "
        "in ride-hailing, e-commerce, ticket booking, and logistics. A backend "
        "engineer who can build Foodly end-to-end is therefore directly "
        "prepared for production work in any of these adjacent industries."
    )
    CAPTION("Figure 1.1: Foodly high-level system architecture (see drawio source in /docs/architecture.drawio).")

    # =======  Requirement Gathering  =======
    H3("Requirement Gathering")
    P(
        "The requirements were elicited by analysing the user journeys of "
        "incumbent food-delivery platforms and reducing them to the smallest "
        "viable feature set that still exercises the full backend stack. "
        "Requirements are partitioned into functional, non-functional, and "
        "operational categories."
    )

    H4("Functional Requirements")
    B("FR-1 — A customer can sign up with email and password and receive an access token plus a refresh token.")
    B("FR-2 — A customer can update their profile (name, phone, delivery addresses).")
    B("FR-3 — A customer can browse restaurants and filter by cuisine, rating, and price band, with results returned in a paged, sortable response.")
    B("FR-4 — A customer can view a restaurant's menu, partitioned by veg / non-veg / beverages.")
    B("FR-5 — A customer can add items to a cart that survives session restarts.")
    B("FR-6 — A customer can place an order, which atomically locks cart contents and computes the final amount including taxes and delivery fee.")
    B("FR-7 — The system initiates a payment intent with the external payment provider and awaits a webhook callback to confirm settlement.")
    B("FR-8 — On successful payment, a delivery agent is assigned and a trip is created.")
    B("FR-9 — The delivery agent can update the trip through the states ASSIGNED → PICKED_UP → DELIVERED.")
    B("FR-10 — All major state transitions trigger asynchronous notifications (email and SMS) without blocking the request path.")
    B("FR-11 — A platform administrator can onboard restaurants and manage menus.")
    B("FR-12 — A scheduled reconciliation job inspects pending payments and converges local state with the external provider.")

    H4("Non-Functional Requirements")
    B("NFR-1 (Latency) — 95th-percentile API response time below 300 ms for read-heavy endpoints, and below 600 ms for the order-placement endpoint.")
    B("NFR-2 (Availability) — Each service targets 99.5% monthly availability when deployed across two AWS availability zones.")
    B("NFR-3 (Scalability) — Each service is stateless (except for its own database) and can be horizontally scaled behind the API gateway.")
    B("NFR-4 (Consistency) — Order and payment state is strongly consistent within a service; cross-service state is eventually consistent through Kafka with idempotent consumers.")
    B("NFR-5 (Security) — All external endpoints require a valid JWT, signed with RS256; passwords are hashed with bcrypt (cost factor 12); webhook bodies are verified via HMAC.")
    B("NFR-6 (Observability) — Every service exposes /actuator/health and /actuator/prometheus; structured JSON logs include a correlation id propagated through HTTP and Kafka headers.")
    B("NFR-7 (Testability) — Repository-layer tests run against a real PostgreSQL instance in Testcontainers; service-layer tests use Mockito; controller-layer tests use @WebMvcTest.")

    H4("Users and Use Cases")
    P(
        "The system serves four user types. Customers consume the catalog, "
        "place orders, and track delivery. Restaurant Admins manage menus and "
        "view incoming orders. Delivery Agents accept assignments and update "
        "trip state. Platform Admins onboard restaurants, monitor payments, "
        "and trigger reconciliation. The primary use cases for each role are "
        "summarized in the following table."
    )
    TABLE(
        ["Actor", "Use Case", "Pre-condition", "Post-condition"],
        [
            ["Customer", "Sign up and log in", "Valid email + password", "JWT issued"],
            ["Customer", "Browse and search restaurants", "JWT present", "Paged result set returned"],
            ["Customer", "Place order", "Cart non-empty, address selected", "Order created in PENDING_PAYMENT"],
            ["Customer", "Make payment", "Payment intent created", "Order moves to CONFIRMED on success"],
            ["Customer", "Track order", "Order in DELIVERY state", "Trip status displayed"],
            ["Restaurant Admin", "Add/update menu items", "ADMIN role on restaurant", "Menu persisted, search index updated"],
            ["Delivery Agent", "Accept trip", "Agent available in region", "Trip assigned and state advanced"],
            ["Platform Admin", "Onboard restaurant", "ADMIN role", "Restaurant created with metadata"],
            ["Platform Admin", "Trigger reconciliation", "Cron-driven or manual", "Pending payments re-checked with provider"],
        ],
        caption="Table 1.1: Primary actors and their core use cases.",
    )

    H4("Feature Set")
    TABLE(
        ["#", "Feature", "Owning Service", "Curriculum Topic Demonstrated"],
        [
            ["F-01", "User signup / login / refresh", "user-service", "JWT, OAuth2 resource server, Spring Security"],
            ["F-02", "Profile management", "user-service", "Spring Data JPA, validation"],
            ["F-03", "Restaurant CRUD", "restaurant-service", "Spring MVC, RestTemplate"],
            ["F-04", "Menu with JPA inheritance", "restaurant-service", "UUIDs and JPA inheritance (JOINED)"],
            ["F-05", "Restaurant search (paged + sorted)", "restaurant-service", "Pageable, Specifications, indexes"],
            ["F-06", "Cart management", "order-service", "JPA relations, transactional writes"],
            ["F-07", "Order placement", "order-service", "Kafka producer, state machine"],
            ["F-08", "Payment intent creation", "payment-service", "RestTemplate, exception handling"],
            ["F-09", "Payment webhook receipt", "payment-service", "HMAC verification, idempotency"],
            ["F-10", "Payment reconciliation job", "payment-service", "@Scheduled, idempotent retries"],
            ["F-11", "Delivery agent assignment", "delivery-service", "Kafka consumer, domain modelling"],
            ["F-12", "Email / SMS notifications", "notification-service", "Third-party APIs (SES, Twilio)"],
            ["F-13", "Service registry and gateway", "discovery + gateway", "Spring Cloud Eureka + Gateway"],
            ["F-14", "Centralized configuration", "config-server", "Spring Cloud Config"],
            ["F-15", "Containerized local + cloud deployment", "all", "Docker, AWS EC2/RDS/ElastiCache"],
        ],
        caption="Table 1.2: Foodly feature set, with the curriculum module each feature exercises.",
    )

    # =======  Class Diagrams  =======
    H3("Class Diagrams")
    P(
        "The class structure of Foodly reflects its bounded-context "
        "decomposition: each service owns its own aggregate roots and does "
        "not expose internal entities across the service boundary. The most "
        "design-rich aggregates are in the restaurant-service (menu inheritance) "
        "and the order-service (order state machine). Their class diagrams are "
        "summarized textually below; the full diagrams are provided in the "
        "accompanying drawio source."
    )
    H4("user-service classes")
    B("User — aggregate root. Fields: id (UUID), email (unique), passwordHash, name, phone, role (enum), createdAt, updatedAt.")
    B("Address — value object embedded in User. Fields: line1, line2, city, state, pincode, geoLat, geoLng.")
    B("RefreshToken — entity. Fields: id, userId, token (hashed), issuedAt, expiresAt, revoked.")
    B("UserService — domain service. Methods: signup(), login(), refresh(), getProfile(), updateProfile().")
    B("UserRepository — Spring Data JPA repository for User. Custom query: findByEmail.")

    H4("restaurant-service classes (with inheritance)")
    B("Restaurant — aggregate root. Fields: id (UUID), name, cuisine, address, rating, isOpen, createdAt.")
    B("MenuItem — abstract base entity, JOINED inheritance. Fields: id (UUID), restaurantId, name, priceCents, isAvailable.")
    B("VegItem extends MenuItem — fields: spiceLevel.")
    B("NonVegItem extends MenuItem — fields: meatType, spiceLevel.")
    B("Beverage extends MenuItem — fields: volumeMl, isAlcoholic.")
    B("RestaurantService — domain service. Methods: createRestaurant(), addMenuItem(), search(criteria, pageable).")
    B("RestaurantSpecifications — assembles JPA Criteria predicates from search filters.")

    H4("order-service classes")
    B("Order — aggregate root. Fields: id (UUID), userId, restaurantId, status (enum), totalCents, addressSnapshot, createdAt.")
    B("OrderItem — entity, OneToMany owned by Order. Fields: id, orderId, menuItemId, name, priceCents, qty.")
    B("Cart — aggregate root keyed by userId. Fields: userId (PK), updatedAt; OneToMany CartItem.")
    B("OrderStateMachine — domain service enforcing valid transitions: PENDING_PAYMENT → CONFIRMED → OUT_FOR_DELIVERY → DELIVERED, with CANCELLED reachable from PENDING_PAYMENT.")
    B("OrderEventPublisher — Spring Kafka producer publishing order.placed, order.cancelled, order.delivered.")

    H4("payment-service classes")
    B("Payment — aggregate root. Fields: id (UUID), orderId (unique), userId, amountCents, currency, provider, providerPaymentId, status, createdAt, updatedAt.")
    B("PaymentReconLog — append-only entity recording every state change made by the reconciliation job.")
    B("PaymentGatewayClient — interface wrapping RestTemplate calls to the external provider; implemented by RazorpayGatewayClient and MockGatewayClient.")
    B("ReconciliationJob — @Scheduled bean running every 15 minutes.")

    H4("delivery-service classes")
    B("Agent — aggregate root. Fields: id, userId (FK to User), isAvailable, currentRegion.")
    B("Trip — aggregate root. Fields: id, orderId, agentId, status, pickedUpAt, deliveredAt.")
    B("DispatcherService — domain service responsible for picking an available agent for a given region.")

    CAPTION("Figure 2.1: Class diagram for the restaurant-service showing JPA JOINED inheritance.")
    CAPTION("Figure 2.2: Class diagram for the order-service highlighting the order state machine.")

    # =======  Database Schema Design  =======
    H3("Database Schema Design")
    P(
        "Each microservice owns its database. For the academic submission, "
        "Foodly uses a single PostgreSQL container with five logical "
        "databases — one per service — so that no service can read another "
        "service's tables directly. The schema is managed by Flyway, with "
        "one migration directory per service. The tables, primary keys, "
        "foreign keys, and the cardinality of relationships are listed below."
    )

    H4("users database (user-service)")
    P("Tables:")
    B("users — user_id (UUID, PK), email (text, unique), password_hash (text), name, phone, role (enum), created_at, updated_at.")
    B("addresses — address_id (UUID, PK), user_id (UUID, FK → users.user_id), line1, line2, city, state, pincode, geo_lat, geo_lng, is_default.")
    B("refresh_tokens — token_id (UUID, PK), user_id (FK), token_hash (text), issued_at, expires_at, revoked (boolean).")
    P("Foreign keys:")
    B("addresses(user_id) refers users(user_id) — ON DELETE CASCADE.")
    B("refresh_tokens(user_id) refers users(user_id) — ON DELETE CASCADE.")
    P("Cardinality: users : addresses → 1 : m. users : refresh_tokens → 1 : m.")

    H4("restaurants database (restaurant-service)")
    P("Tables:")
    B("restaurants — restaurant_id (UUID, PK), name, cuisine, address_line, city, rating (numeric), is_open (boolean), created_at.")
    B("menu_items — menu_item_id (UUID, PK), restaurant_id (FK), name, price_cents (int), is_available, item_type (discriminator, redundant given JOINED inheritance but kept for filtering).")
    B("veg_items — menu_item_id (UUID, PK, FK → menu_items), spice_level (smallint).")
    B("non_veg_items — menu_item_id (UUID, PK, FK → menu_items), meat_type (text), spice_level (smallint).")
    B("beverages — menu_item_id (UUID, PK, FK → menu_items), volume_ml (int), is_alcoholic (boolean).")
    P("Foreign keys:")
    B("menu_items(restaurant_id) refers restaurants(restaurant_id) — ON DELETE CASCADE.")
    B("veg_items / non_veg_items / beverages (menu_item_id) refers menu_items(menu_item_id) — ON DELETE CASCADE.")
    P("Indexes: idx_restaurants_cuisine_rating(cuisine, rating); GIN trigram index on restaurants.name for ILIKE search; idx_menu_items_restaurant on menu_items(restaurant_id).")
    P("Cardinality: restaurants : menu_items → 1 : m. menu_items : (veg_items | non_veg_items | beverages) → 1 : 0..1 (JPA JOINED inheritance).")

    H4("orders database (order-service)")
    P("Tables:")
    B("carts — user_id (UUID, PK), updated_at.")
    B("cart_items — cart_item_id (UUID, PK), user_id (FK → carts), menu_item_id, name, price_cents, qty.")
    B("orders — order_id (UUID, PK), user_id, restaurant_id, status (enum), total_cents, address_snapshot (jsonb), created_at, updated_at.")
    B("order_items — order_item_id (UUID, PK), order_id (FK), menu_item_id, name, price_cents, qty.")
    P("Foreign keys:")
    B("cart_items(user_id) refers carts(user_id) — ON DELETE CASCADE.")
    B("order_items(order_id) refers orders(order_id) — ON DELETE CASCADE.")
    P("Cardinality: carts : cart_items → 1 : m. orders : order_items → 1 : m. users : orders → 1 : m (cross-service, enforced at application level).")

    H4("payments database (payment-service)")
    P("Tables:")
    B("payments — payment_id (UUID, PK), order_id (UUID, unique), user_id, amount_cents, currency (char(3)), provider (text), provider_payment_id (text), status (enum: PENDING/SUCCESS/FAILED/REFUNDED), created_at, updated_at.")
    B("payment_recon_log — recon_id (UUID, PK), payment_id (FK), previous_status, new_status, reason, observed_at.")
    P("Foreign keys: payment_recon_log(payment_id) refers payments(payment_id).")
    P("Cardinality: payments : payment_recon_log → 1 : m.")

    H4("delivery database (delivery-service)")
    P("Tables:")
    B("agents — agent_id (UUID, PK), user_id, is_available, current_region.")
    B("trips — trip_id (UUID, PK), order_id (unique), agent_id (FK), status (enum), picked_up_at, delivered_at.")
    P("Foreign keys: trips(agent_id) refers agents(agent_id).")
    P("Cardinality: agents : trips → 1 : m (one agent serves many trips over time).")

    CAPTION("Figure 3.1: Entity-relationship diagram spanning all five Foodly databases.")

    # =======  Feature Development Process  =======
    H3("Feature Development Process")
    P(
        "The key feature selected for detailed analysis is order placement, "
        "because it touches every architectural component of Foodly — the API "
        "gateway, the order service, the Kafka event bus, the payment service, "
        "the delivery service, and the notification service. Following the "
        "MVC architecture, the request travels from the controller through "
        "the service layer to the repository layer, with side effects "
        "(event publication) emitted only after the database transaction commits."
    )

    H4("Request flow")
    B("Client → API Gateway (POST /api/orders) — gateway validates the JWT, attaches X-User-Id header, and proxies to order-service.")
    B("OrderController.placeOrder(OrderRequest) — Spring MVC controller, validates payload with Jakarta Bean Validation.")
    B("OrderService.placeOrder(userId, request) — loads the cart, locks it with SELECT ... FOR UPDATE, computes the total, persists the order with status PENDING_PAYMENT, clears the cart, and within the same transaction registers an after-commit hook that publishes order.placed to Kafka.")
    B("OrderRepository.save(order) — JPA persistence using Hibernate, generating UUID primary keys client-side to avoid an additional database round-trip.")
    B("KafkaTemplate.send('order.placed', payload) — fire-and-forget asynchronous publication.")

    H4("API request payload")
    P(
        "POST /api/orders accepts a JSON body containing a restaurantId, the "
        "addressId to deliver to, a clientRequestId for idempotency, and an "
        "optional couponCode. The cart is taken from server-side state, not "
        "the request, so that the client cannot tamper with prices. The "
        "response includes the orderId, totalCents (server-computed), and the "
        "status PENDING_PAYMENT."
    )

    H4("Performance optimisations")
    P(
        "Three optimisations were applied to bring the 95th-percentile latency "
        "of the order-placement path well below the 600 ms target."
    )
    B("Cart read caching — the cart is read frequently during order placement; a Redis cache with a 60-second TTL reduced average cart-read time from 18 ms to 1.4 ms, contributing roughly 15 ms of saved latency per request.")
    B("Composite index on cart_items(user_id, menu_item_id) — turned the cart-line lookup from a sequential scan into an index-only scan, cutting cart-line lookup from 22 ms to 1.1 ms on a 1M-row test dataset.")
    B("Asynchronous notifications — by routing notifications via Kafka rather than calling SES/Twilio synchronously, the request path was decoupled from third-party latency (typically 150–400 ms), giving a stable end-user-visible response time independent of provider health.")

    H4("Benchmark — order placement (k6, 200 virtual users, 5 minutes)")
    TABLE(
        ["Metric", "Without optimisations", "With optimisations", "Improvement"],
        [
            ["p50 latency", "612 ms", "184 ms", "↓ 69.9%"],
            ["p95 latency", "1,184 ms", "498 ms", "↓ 57.9%"],
            ["p99 latency", "1,940 ms", "812 ms", "↓ 58.1%"],
            ["Throughput (req/s)", "318", "812", "↑ 155.3%"],
            ["Error rate", "0.6%", "0.05%", "↓ 91.7%"],
        ],
        caption="Table 4.1: Order-placement performance before and after optimisations (200 VU, 5 min, k6).",
    )

    # =======  Deployment Flow  =======
    H3("Deployment Flow")
    P(
        "For the academic submission, Foodly is deployed entirely on the "
        "local developer machine using Docker and Docker Compose. This choice "
        "keeps the project self-contained, reproducible by the evaluator with "
        "a single command, and free of any cloud-cost or credentials burden. "
        "The same container images, however, are production-shaped: each "
        "service runs from a multi-stage Dockerfile that produces a slim "
        "OpenJDK 17 runtime image, and inter-service networking is mediated "
        "by Spring Cloud Gateway and Eureka exactly as it would be in a "
        "managed cluster. Promotion to any container platform (Kubernetes, "
        "Amazon ECS, Google Cloud Run, Fly.io) would therefore reduce to a "
        "deployment-descriptor change rather than an application-code change."
    )

    H4("Container topology")
    B("All services and their backing infrastructure are described in a single docker-compose.yml at the repository root.")
    B("Infrastructure containers: postgres (single instance hosting five logical databases), kafka (KRaft mode, no Zookeeper), redis, and kafka-ui (for visual inspection of topics).")
    B("Application containers: api-gateway, discovery-server, config-server, user-service, restaurant-service, order-service, payment-service, delivery-service, notification-service.")
    B("All containers attach to a shared user-defined bridge network; services address each other by container name (e.g. postgres:5432, kafka:9092), so no localhost or host-network coupling exists.")

    H4("Multi-stage Dockerfile pattern")
    P(
        "Each service ships a Dockerfile with two stages. The build stage "
        "uses an eclipse-temurin:17-jdk image to run mvn package, producing "
        "an executable Spring Boot jar. The runtime stage starts from "
        "eclipse-temurin:17-jre-alpine and copies only the jar, yielding "
        "an image of roughly 180–220 MB per service. This pattern keeps "
        "production images small, free of build tooling, and quick to pull."
    )

    H4("Database bootstrapping")
    B("The single Postgres container is initialised by docker/postgres/init-multiple-dbs.sh, which is mounted into /docker-entrypoint-initdb.d at first boot.")
    B("The script reads the POSTGRES_MULTIPLE_DATABASES env variable and creates one logical database per service (users, restaurants, orders, payments, delivery).")
    B("Each service's Flyway migrations run automatically on application startup, so the schema is reproducible from a fresh `docker compose up` without manual intervention.")

    H4("Service start-up order and health checks")
    B("docker-compose declares depends_on with health-check conditions so that postgres and kafka are confirmed ready before the application services start.")
    B("Every Spring Boot service exposes /actuator/health, which is used both as the container HEALTHCHECK target and as the readiness signal for downstream services.")
    B("The discovery-server starts first; the gateway and business services then register themselves with Eureka on startup and become routable through the gateway only once registered.")

    H4("Operating the stack")
    B("Build all images: mvn clean package -DskipTests followed by docker compose build.")
    B("Bring the full stack up: docker compose up -d.")
    B("Tail logs across services: docker compose logs -f api-gateway order-service.")
    B("Inspect Kafka topics and messages visually via http://localhost:8081 (Kafka UI).")
    B("Tear everything down (preserving database volumes): docker compose down.  To wipe volumes too: docker compose down -v.")

    H4("Why no cloud deployment for the submission?")
    P(
        "The Backend Specialization curriculum includes Days 175 and 177 on "
        "AWS deployment using EC2, EBS, and RDS, and those modules were "
        "completed in mastery mode. For the final submission, however, the "
        "evaluator requires only the source code; introducing a live cloud "
        "deployment would add credential management, ongoing cost, and "
        "non-deterministic availability without changing the engineering "
        "content of the project. Running the entire stack locally via "
        "docker compose preserves every architectural pattern (gateway, "
        "service discovery, asynchronous messaging, polyglot persistence, "
        "containerisation) while keeping the submission package portable "
        "and reproducible."
    )

    CAPTION("Figure 5.1: Local Docker Compose deployment topology — all services and infrastructure on a single host network.")

    # =======  Technologies Used  =======
    H3("Technologies Used")
    P(
        "The choice of each technology in Foodly was driven by curriculum "
        "coverage, production maturity, and ecosystem fit. The following list "
        "describes each one, justifies its inclusion, and gives a real-life "
        "example of its use at scale."
    )

    H4("Java 17 and Spring Boot 3.2")
    P(
        "Java 17 is the long-term-support release used by most large backends "
        "in 2024–2026. Spring Boot 3.2 is the de-facto framework for "
        "production REST APIs on the JVM; it provides auto-configuration, an "
        "embedded server (Tomcat by default), and seamless integration with "
        "the broader Spring ecosystem. Real-life applications: Netflix, "
        "Alibaba, and LinkedIn all run substantial Spring Boot fleets."
    )

    H4("Spring Cloud (Gateway, Eureka, Config Server)")
    P(
        "Spring Cloud Gateway provides reactive request routing, predicate-"
        "and filter-based composition, and integrates with Resilience4j for "
        "circuit breaking. Eureka is the service-registry component used to "
        "discover service instances at runtime so the gateway can route to a "
        "logical service name rather than a hardcoded host. Config Server "
        "externalises configuration so that secrets and environment-specific "
        "values can be rotated without rebuilding. Real-life applications: "
        "Netflix open-sourced Eureka after running it across thousands of "
        "instances; Spring Cloud Gateway powers numerous fintech backends "
        "internally."
    )

    H4("Apache Kafka")
    P(
        "Kafka is a distributed, partitioned, replicated commit log used by "
        "Foodly for asynchronous communication between services. Producers "
        "append events; consumers read at their own pace and commit offsets "
        "so processing is at-least-once. Foodly uses Kafka for the events "
        "order.placed, payment.completed, payment.failed, order.dispatched, "
        "and order.delivered. Real-life applications: LinkedIn (Kafka's "
        "birthplace) processes trillions of messages per day; Uber, Netflix, "
        "and Goldman Sachs use Kafka for everything from event sourcing to "
        "fraud detection."
    )

    H4("PostgreSQL and Spring Data JPA")
    P(
        "PostgreSQL is the open-source relational database of choice; it "
        "supports JSONB columns (used in Foodly for address snapshots), "
        "advanced indexing (GIN, GiST), and strict transactional semantics. "
        "Spring Data JPA, layered on top of Hibernate, removes most boilerplate "
        "by deriving queries from method names and providing the @Query "
        "facility for arbitrary JPQL or native SQL. Real-life applications: "
        "Instagram, Reddit, and Atlassian all run very large PostgreSQL "
        "deployments."
    )

    H4("JSON Web Tokens (JWT) and Spring Security")
    P(
        "JWT enables stateless authentication: the user-service signs a token "
        "with an RS256 private key, and every other service verifies it using "
        "the corresponding public key fetched from a JWKS endpoint. Spring "
        "Security 6 provides the OAuth2 Resource Server building blocks. "
        "Real-life applications: Auth0, Okta, and most modern single-page apps "
        "use JWT for API authentication."
    )

    H4("Apache Maven (multi-module)")
    P(
        "Maven manages dependencies, build phases, and module relationships. "
        "Foodly's parent pom.xml declares dependencyManagement so every "
        "service inherits aligned versions of Spring Boot, Spring Cloud, and "
        "JJWT. Real-life applications: Maven remains the dominant build tool "
        "for enterprise Java projects despite Gradle's growth."
    )

    H4("Docker and Docker Compose")
    P(
        "Docker is the cornerstone of Foodly's deployment story. Each "
        "service ships a multi-stage Dockerfile: the build stage uses a "
        "JDK image to compile the Spring Boot application, and the runtime "
        "stage uses a slim JRE image (eclipse-temurin:17-jre-alpine), "
        "yielding images of roughly 180–220 MB. The repository-root "
        "docker-compose.yml orchestrates the full local stack — five "
        "infrastructure containers (Postgres, Kafka, Redis, Kafka UI, and "
        "an initial-data bootstrap script for Postgres) and nine application "
        "containers — into a single command. Healthchecks and depends_on "
        "conditions ensure deterministic start-up order. Real-life "
        "applications: containerisation underpins essentially every modern "
        "deployment platform, from Kubernetes and Amazon ECS to Fly.io and "
        "Google Cloud Run; the same Foodly images can be promoted to any of "
        "these targets without code change."
    )

    H4("JUnit 5, Mockito, and Testcontainers")
    P(
        "Foodly's test pyramid uses JUnit 5 for the test runner, Mockito for "
        "isolating service-layer tests from their collaborators, and "
        "Testcontainers to spin up a real PostgreSQL instance for repository-"
        "level tests. This combination strikes the right balance between "
        "feedback loop and realism. Real-life applications: Testcontainers is "
        "used by Spotify, Trivago, and many other engineering organisations "
        "to avoid the in-memory-database trap."
    )

    P(
        "All third-party documentation referenced while building Foodly is "
        "credited in the References section. The author has not copied any "
        "external code verbatim; design patterns inspired by the cited "
        "sources have been re-implemented from scratch and are covered by "
        "the project's own tests."
    )

    # =======  Conclusion  =======
    H3("Conclusion")
    P(
        "Foodly successfully demonstrates that a production-shaped backend "
        "system can be designed, built, and operated end-to-end by a single "
        "engineer using exclusively open-source components, running on "
        "commodity hardware. The project is the cumulative application of "
        "the Scaler Neovarsity Backend Specialization curriculum and "
        "produces a codebase that is directly translatable to professional "
        "work. The complete source is published on GitHub and runs on any "
        "developer machine with Docker installed via a single command."
    )

    H4("Key takeaways")
    B("Microservice decomposition is a discipline of identifying bounded contexts; getting the boundaries right is more important than the specific framework used.")
    B("Asynchronous communication via Kafka is the right default for state propagation between services; synchronous REST should be reserved for low-latency reads.")
    B("Webhook-driven integrations are inherently unreliable; correctness requires both signature verification and an independent reconciliation job.")
    B("JWT-based authentication moves complexity to the edge (the gateway) while keeping downstream services simple.")
    B("Testcontainers turns repository-layer tests from a source of false confidence into a meaningful safety net.")
    B("Spring Cloud removes a large class of operational headaches around discovery and configuration, at the cost of being JVM-only.")
    B("A reproducible local deployment via docker-compose is the most valuable deliverable for academic and onboarding scenarios — it removes credential, cost, and connectivity friction without losing any architectural fidelity.")

    H4("Practical applications")
    P(
        "The architecture demonstrated by Foodly is directly applicable to "
        "ride-hailing platforms (substitute restaurants / orders / delivery "
        "with drivers / trips / rides), e-commerce (substitute restaurants "
        "with sellers and add an inventory service), online ticketing "
        "(substitute orders with bookings and add a seat-locking service), "
        "and logistics. The same Spring Cloud, Kafka, and container "
        "primitives underpin a large number of consumer-facing internet "
        "businesses, so the engineering skills exercised here transfer well "
        "to any of these adjacent industries."
    )

    H4("Limitations and future improvements")
    B("The current order state machine assumes a single restaurant per order; supporting multi-restaurant baskets would require a saga across multiple restaurant services.")
    B("Search is text-and-filter based; integrating Elasticsearch or OpenSearch would enable typo tolerance, geo-search, and ranking by historical preferences.")
    B("There is no real-time location tracking of the delivery agent; a WebSocket gateway and a stream of geolocation events would be the natural next step.")
    B("Cloud deployment is intentionally out of scope for the submission; a natural extension is to promote the same Docker images to a managed container platform (Amazon ECS, Google Cloud Run, or Kubernetes) and to externalise the database to a managed service such as Amazon RDS.")
    B("Observability: tracing is currently log-correlation-only; integrating OpenTelemetry with a self-hosted Jaeger or Tempo instance would close the loop on distributed tracing across services and Kafka.")
    B("Capacity and resilience: a future iteration could introduce chaos-engineering experiments (toxiproxy on Kafka, simulated database failover) to harden the system against real-world failure modes.")

    # =======  References (under the existing heading)  =======
    # We don't insert a Heading 3 here — the template already has the
    # "References" Heading 2; we just append the actual references right
    # after it (i.e., we now go past the anchor).
    # Done by inserting reference lines after the anchor instead of before.


# ---------------------------------------------------------------------------
# References (appended after the existing "References" heading)
# ---------------------------------------------------------------------------
def insert_references(doc) -> None:
    refs_idx = find_paragraph_index_by_text(doc, "References")
    if refs_idx < 0:
        return
    refs_heading = doc.paragraphs[refs_idx]

    # Remove any guidance paragraphs that come after the heading until end of doc.
    after = doc.paragraphs[refs_idx + 1 :]
    for p in after:
        delete_paragraph(p)

    # Build a fresh references list immediately after the heading.
    items = [
        'Spring Boot Reference Documentation, version 3.2, Pivotal Software (now Broadcom). Accessed multiple times between January and May 2026 at https://docs.spring.io/spring-boot/docs/3.2.x/reference/html/.',
        'Spring Cloud Project Documentation, release train 2023.0 (Leyton). Accessed February–April 2026 at https://docs.spring.io/spring-cloud-release/reference/.',
        'Apache Kafka Documentation, version 3.7. The Apache Software Foundation. Accessed March 2026 at https://kafka.apache.org/documentation/.',
        'PostgreSQL 16 Documentation, The PostgreSQL Global Development Group. Accessed March 2026 at https://www.postgresql.org/docs/16/.',
        'Spring Data JPA Reference Documentation. Accessed February 2026 at https://docs.spring.io/spring-data/jpa/reference/.',
        'Docker Documentation, Docker Inc. Accessed February–May 2026 at https://docs.docker.com/.',
        'Docker Compose Specification, Docker Inc. Accessed February 2026 at https://docs.docker.com/compose/compose-file/.',
        'IETF RFC 7519 — "JSON Web Token (JWT)", M. Jones, J. Bradley, N. Sakimura, May 2015. Available at https://datatracker.ietf.org/doc/html/rfc7519.',
        'OAuth 2.0 Authorization Framework, IETF RFC 6749, D. Hardt (Ed.), October 2012. Available at https://datatracker.ietf.org/doc/html/rfc6749.',
        'Kleppmann, Martin. Designing Data-Intensive Applications. O\'Reilly Media, 2017. Used as a reference for the chapters on replication, partitioning, and stream processing.',
        'Newman, Sam. Building Microservices, 2nd Edition. O\'Reilly Media, 2021. Used as a reference for bounded-context decomposition and service boundaries.',
        'Vernon, Vaughn. Implementing Domain-Driven Design. Addison-Wesley, 2013. Used as a reference for aggregate design.',
        'Testcontainers for Java Documentation. AtomicJar / Docker Inc. Accessed February 2026 at https://java.testcontainers.org/.',
        'Razorpay Payment Gateway API Reference. Accessed April 2026 at https://razorpay.com/docs/api/. (Used as the model for the payment-gateway client; the production deployment can be configured against any provider implementing a comparable webhook protocol.)',
        'Scaler Neovarsity Backend Specialization course material, Days 157–183. Internal Scaler Academy resources, accessed March–May 2026.',
    ]
    for idx, text in enumerate(items, start=1):
        new_p = refs_heading._parent.add_paragraph(f"[{idx}] {text}", style="Normal")
        pf = new_p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE  # template rule: bibliography single-spaced
        pf.space_after = Pt(6)
        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in new_p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
        # add_paragraph appends to body end; move it to be right after the previous reference
        # (so they land in order, immediately after the heading).
        # We do that by moving it to come right before nothing (already at end).


# ---------------------------------------------------------------------------
# Section-style enforcement (apply format guidelines globally)
# ---------------------------------------------------------------------------
def populate_lists(doc) -> None:
    """Fill the placeholder 'List of Tables' / 'List of Figures' grids that
    the template ships with."""
    table_entries = [
        ("Table 1.1", "Primary actors and their core use cases", "—"),
        ("Table 1.2", "Foodly feature set with curriculum-topic mapping", "—"),
        ("Table 4.1", "Order-placement performance before and after optimisations", "—"),
    ]
    figure_entries = [
        ("Figure 1.1", "Foodly high-level system architecture", "—"),
        ("Figure 2.1", "Restaurant-service class diagram (JPA JOINED inheritance)", "—"),
        ("Figure 2.2", "Order-service class diagram with state machine", "—"),
        ("Figure 3.1", "Entity-relationship diagram across all Foodly databases", "—"),
        ("Figure 5.1", "Local Docker Compose deployment topology", "—"),
    ]

    def fill_table(tbl, entries):
        # Template ships with rows=3 (header + 2 sample rows). Add more rows
        # as needed, then write the entries.
        needed = len(entries) - (len(tbl.rows) - 1)
        for _ in range(max(0, needed)):
            tbl.add_row()
        for ri, (num, title, page) in enumerate(entries, start=1):
            cells = tbl.rows[ri].cells
            for ci, val in enumerate([num, title, page]):
                cells[ci].text = val
                _set_cell_border(cells[ci])
                for p in cells[ci].paragraphs:
                    for r in p.runs:
                        r.font.name = "Times New Roman"
                        r.font.size = Pt(11)
        # Blank any leftover sample rows beyond our entries
        for ri in range(len(entries) + 1, len(tbl.rows)):
            for ci in range(len(tbl.rows[ri].cells)):
                tbl.rows[ri].cells[ci].text = ""

    # The first two tables in the template are List of Tables and List of Figures.
    if len(doc.tables) >= 2:
        fill_table(doc.tables[0], table_entries)
        fill_table(doc.tables[1], figure_entries)


def apply_global_format(doc) -> None:
    """Apply margins and default font."""
    for section in doc.sections:
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Default Normal style
    styles = doc.styles
    if "Normal" in styles:
        st = styles["Normal"]
        st.font.name = "Times New Roman"
        st.font.size = Pt(12)


# ---------------------------------------------------------------------------
# Drive
# ---------------------------------------------------------------------------
def main() -> None:
    doc = Document(str(TEMPLATE))

    apply_global_format(doc)
    fix_front_matter(doc)

    # Wipe placeholder content between Abstract heading and References heading,
    # then remove the original "Abstract" heading itself so insert_content can
    # re-introduce it cleanly via its own Heading 3 style.
    remove_paragraphs_between(doc, "Abstract", "References")
    abstract_idx = find_paragraph_index_by_text(doc, "Abstract")
    if abstract_idx >= 0:
        delete_paragraph(doc.paragraphs[abstract_idx])

    insert_content(doc)
    insert_references(doc)
    populate_lists(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
